import argparse
import pandas as pd
from fpdf import FPDF
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import matplotlib.pyplot as plt
from sqlalchemy import create_engine
from abc import ABC, abstractmethod
from openpyxl import Workbook
from openpyxl.styles import PatternFill
import os
import time

# Diretório de saída para os relatórios
OUTPUT_DIR = r"C:/iAxxMES/OutputRelatorios/"

# Garante que o diretório de saída exista
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Configura a conexão usando SQLAlchemy
def conectar_bd():
    conexao_str = 'mysql+mysqlconnector://root@localhost/iAxxMES'
    engine = create_engine(conexao_str)
    return engine

# Dicionário de cores para o status
STATUS_CORES = {
    'Rodando': '008000',  # Verde
    'Parada': 'FF0000',  # Vermelho
    'Setup': 'ADD8E6',   # Azul claro
    'Carga de fio': 'FFA500',  # Laranja
    'Sem programação': '808080'  # Cinza
}

# Função para calcular o tempo em cada status corretamente
def calcular_tempo_no_status(dados):
    # Faz uma cópia do DataFrame para evitar o SettingWithCopyWarning
    dados = dados.copy()
    dados = dados.sort_values(['maquina_id', 'data_hora'])
    dados['tempo_no_status'] = dados.groupby('maquina_id')['data_hora'].diff().shift(-1)
    dados['tempo_no_status'] = dados['tempo_no_status'].fillna(pd.Timedelta(seconds=0))
    return dados

# Configuração dos argumentos de linha de comando
parser = argparse.ArgumentParser(description="Geração de Relatórios")
parser.add_argument("--tipo_relatorio", required=True, choices=["RPM", "Status", "Eficiência"], help="Tipo de relatório a ser gerado")
parser.add_argument("--maquina_id", type=int, help="ID da máquina para o relatório (opcional para todas as máquinas)")
parser.add_argument("--data_inicio", required=True, help="Data e hora de início no formato AAAA-MM-DD HH:MM:SS")
parser.add_argument("--data_fim", required=True, help="Data e hora de término no formato AAAA-MM-DD HH:MM:SS")
parser.add_argument("--formatos", required=True, nargs="+", choices=["pdf", "excel", "word"], help="Formatos de saída: pdf, excel, word")

args = parser.parse_args()

# Exemplo de como utilizar os argumentos
maquina_id = args.maquina_id
data_inicio = args.data_inicio
data_fim = args.data_fim
tipo_relatorio = args.tipo_relatorio
formatos = args.formatos

# Classe base para relatórios
class Relatorio(ABC):
    def __init__(self, maquina_id=None, data_inicio=None, data_fim=None):
        self.maquina_id = maquina_id
        self.data_inicio = data_inicio
        self.data_fim = data_fim
        self.dados = None
        self.engine = conectar_bd()

    @abstractmethod
    def obter_dados(self):
        pass

    @abstractmethod
    def gerar_grafico(self, dados, titulo):
        pass

    def gerar_pdf(self, titulo):
        if self.maquina_id is not None:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, titulo, 0, 1, 'C')
            self.gerar_grafico(self.dados, f"Máquina {self.maquina_id}")
            chart_path = os.path.join(OUTPUT_DIR, f"{self.__class__.__name__}_chart_maquina_{self.maquina_id}.png")
            pdf.image(chart_path, x=10, y=30, w=180)
            pdf_output_path = os.path.join(OUTPUT_DIR, f"{self.__class__.__name__}_Relatorio.pdf")
            pdf.output(pdf_output_path)
            print(f"PDF gerado em: {pdf_output_path}")
            time.sleep(1)  # Pausa para garantir gravação

    def gerar_excel(self):
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Relatório Todas Máquinas" if self.maquina_id is None else f"Máquina {self.maquina_id}"

        colunas = list(self.dados.columns)
        sheet.append(colunas)

        for idx, row in self.dados.iterrows():
            linha = [row[col] for col in self.dados.columns]
            sheet.append(linha)

            status = row['status'] if 'status' in row else None
            if status in STATUS_CORES:
                cor = STATUS_CORES[status]
                status_cell = sheet.cell(row=sheet.max_row, column=colunas.index('status') + 1)
                status_cell.fill = PatternFill(start_color=cor, end_color=cor, fill_type="solid")

            if self.maquina_id is None and (idx < len(self.dados) - 1) and row['maquina_id'] != self.dados.loc[idx + 1, 'maquina_id']:
                sheet.append([""] * len(colunas))

        excel_output_path = os.path.join(OUTPUT_DIR, f"{self.__class__.__name__}_Relatorio.xlsx")
        workbook.save(excel_output_path)
        print(f"Excel gerado em: {excel_output_path}")
        time.sleep(1)  # Pausa para garantir gravação

    def gerar_word(self):
        doc = Document()
        titulo = f'Relatório de {self.__class__.__name__}' + (f' - Máquina {self.maquina_id}' if self.maquina_id else ' - Todas as Máquinas')
        doc.add_heading(titulo, 0)

        if self.maquina_id is None:
            maquinas = self.dados['maquina_id'].unique()
            for maquina in maquinas:
                dados_maquina = self.dados[self.dados['maquina_id'] == maquina]
                if 'tempo_no_status' in dados_maquina.columns:
                    dados_maquina['tempo_no_status'] = dados_maquina['tempo_no_status'].astype(str)

                doc.add_paragraph(f"Máquina {maquina}", style="Heading 1")
                table = doc.add_table(rows=1, cols=len(dados_maquina.columns))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells

                for idx, col_name in enumerate(dados_maquina.columns):
                    hdr_cells[idx].text = col_name

                for _, row in dados_maquina.iterrows():
                    row_cells = table.add_row().cells
                    for idx, value in enumerate(row):
                        cell = row_cells[idx]
                        cell.text = str(value)
                        if dados_maquina.columns[idx] == 'status' and value in STATUS_CORES:
                            cor_hex = STATUS_CORES[value]
                            cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), cor_hex)))

                doc.add_paragraph()
        else:
            dados_maquina = self.dados
            if 'tempo_no_status' in dados_maquina.columns:
                dados_maquina['tempo_no_status'] = dados_maquina['tempo_no_status'].astype(str)

            table = doc.add_table(rows=1, cols=len(dados_maquina.columns))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for idx, col_name in enumerate(dados_maquina.columns):
                hdr_cells[idx].text = col_name

            for _, row in dados_maquina.iterrows():
                row_cells = table.add_row().cells
                for idx, value in enumerate(row):
                    cell = row_cells[idx]
                    cell.text = str(value)
                    if self.dados.columns[idx] == 'status' and value in STATUS_CORES:
                        cor_hex = STATUS_CORES[value]
                        cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), cor_hex)))

        word_output_path = os.path.join(OUTPUT_DIR, f"{self.__class__.__name__}_Relatorio.docx")
        doc.save(word_output_path)
        print(f"Word gerado em: {word_output_path}")
        time.sleep(1)  # Pausa para garantir gravação

    def gerar_relatorios(self, tipos_relatorio):
        self.obter_dados()
        titulo = f"Relatório de {self.__class__.__name__}" + (f" - Máquina {self.maquina_id}" if self.maquina_id else " - Todas as Máquinas")

        if 'pdf' in tipos_relatorio and self.maquina_id is not None:
            self.gerar_pdf(titulo)

        if 'excel' in tipos_relatorio:
            self.gerar_excel()

        if 'word' in tipos_relatorio:
            self.gerar_word()

# Classe RelatorioRPM
class RelatorioRPM(Relatorio):
    def obter_dados(self):
        base_query = """
            SELECT maquina_id, data_hora, rpm FROM maquina_dados
            WHERE data_hora BETWEEN %(data_inicio)s AND %(data_fim)s
        """
        params = {'data_inicio': self.data_inicio, 'data_fim': self.data_fim}
        if self.maquina_id is not None:
            base_query += " AND maquina_id = %(maquina_id)s"
            params['maquina_id'] = self.maquina_id
        base_query += " ORDER BY maquina_id, data_hora;"
        self.dados = pd.read_sql(base_query, self.engine, params=params)

    def gerar_grafico(self, dados, titulo):
        plt.figure(figsize=(10, 5))
        plt.plot(dados['data_hora'], dados['rpm'], marker='o', linestyle='-', color='b')
        plt.title(titulo)
        plt.xlabel('Data e Hora')
        plt.ylabel('RPM')
        plt.xticks(rotation=45)
        plt.tight_layout()
        chart_path = os.path.join(OUTPUT_DIR, f"{self.__class__.__name__}_chart_maquina_{self.maquina_id}.png")
        plt.savefig(chart_path)
        plt.close()

# Classe RelatorioStatus
class RelatorioStatus(Relatorio):
    def obter_dados(self):
        base_query = """
            SELECT md.maquina_id, md.data_hora, ms.descricao AS status
            FROM maquina_dados md
            JOIN maquina_status ms ON md.status = ms.id
            WHERE md.data_hora BETWEEN %(data_inicio)s AND %(data_fim)s
        """
        params = {'data_inicio': self.data_inicio, 'data_fim': self.data_fim}
        if self.maquina_id is not None:
            base_query += " AND md.maquina_id = %(maquina_id)s"
            params['maquina_id'] = self.maquina_id
        base_query += " ORDER BY md.maquina_id, md.data_hora;"
        self.dados = pd.read_sql(base_query, self.engine, params=params)
        self.dados = calcular_tempo_no_status(self.dados)

    def gerar_grafico(self, dados, titulo):
        plt.figure(figsize=(10, 5))
        status_values, status_labels = pd.factorize(dados['status'])
        plt.plot(dados['data_hora'], status_values, marker='o', linestyle='-', color='g')
        plt.title(titulo)
        plt.xlabel('Data e Hora')
        plt.ylabel('Status')
        plt.xticks(rotation=45)
        plt.yticks(ticks=range(len(status_labels)), labels=status_labels)
        plt.tight_layout()
        chart_path = os.path.join(OUTPUT_DIR, f"{self.__class__.__name__}_chart_maquina_{self.maquina_id}.png")
        plt.savefig(chart_path)
        plt.close()

# Classe RelatorioEficiencia
class RelatorioEficiencia(Relatorio):
    def obter_dados(self):
        base_query = """
            SELECT md.maquina_id, md.data_hora, ms.descricao AS status
            FROM maquina_dados md
            JOIN maquina_status ms ON md.status = ms.id
            WHERE md.data_hora BETWEEN %(data_inicio)s AND %(data_fim)s
        """
        params = {'data_inicio': self.data_inicio, 'data_fim': self.data_fim}
        if self.maquina_id is not None:
            base_query += " AND md.maquina_id = %(maquina_id)s"
            params['maquina_id'] = self.maquina_id
        base_query += " ORDER BY md.maquina_id, md.data_hora;"
        self.dados = pd.read_sql(base_query, self.engine, params=params)
        self.dados = calcular_tempo_no_status(self.dados)

        # Cálculos para o relatório de eficiência
        self.tempo_disponivel = self.dados.loc[~self.dados['status'].isin(['Setup', 'Carga de fio']), 'tempo_no_status'].sum()
        self.tempo_rodando = self.dados.loc[self.dados['status'] == 'Rodando', 'tempo_no_status'].sum()
        self.tempo_parada = self.dados.loc[self.dados['status'].isin(['Parada', 'Sem programação']), 'tempo_no_status'].sum()
        self.tempo_indisponivel = self.dados.loc[self.dados['status'].isin(['Setup', 'Carga de fio']), 'tempo_no_status'].sum()

    def gerar_grafico(self, dados, titulo):
        pass  # Implementação desnecessária para Relatório de Eficiência

    def gerar_excel(self):
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = f"Eficiência Máquina {self.maquina_id}" if self.maquina_id else "Eficiência Todas Máquinas"

        headers = ["Máquina ID", "Tempo Disponível", "Tempo Rodando", "Tempo Parada", "Tempo Indisponível"]
        sheet.append(headers)

        row_data = [
            str(self.maquina_id if self.maquina_id else "Todas as Máquinas"),
            str(self.tempo_disponivel),
            str(self.tempo_rodando),
            str(self.tempo_parada),
            str(self.tempo_indisponivel)
        ]
        sheet.append(row_data)

        excel_output_path = os.path.join(OUTPUT_DIR, f"{self.__class__.__name__}_Relatorio.xlsx")
        workbook.save(excel_output_path)

    def gerar_word(self):
        doc = Document()
        titulo = f'Relatório de Eficiência - Máquina {self.maquina_id}' if self.maquina_id else 'Relatório de Eficiência - Todas as Máquinas'
        doc.add_heading(titulo, 0)

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ["Máquina ID", "Tempo Disponível", "Tempo Rodando", "Tempo Parada", "Tempo Indisponível"]
        for idx, header in enumerate(headers):
            hdr_cells[idx].text = header

        row_cells = table.add_row().cells
        row_cells[0].text = str(self.maquina_id if self.maquina_id else "Todas as Máquinas")
        row_cells[1].text = str(self.tempo_disponivel)
        row_cells[2].text = str(self.tempo_rodando)
        row_cells[3].text = str(self.tempo_parada)
        row_cells[4].text = str(self.tempo_indisponivel)

        word_output_path = os.path.join(OUTPUT_DIR, f"{self.__class__.__name__}_Relatorio.docx")
        doc.save(word_output_path)

# Seleção do tipo de relatório
if tipo_relatorio == "RPM":
    relatorio = RelatorioRPM(maquina_id, data_inicio, data_fim)
elif tipo_relatorio == "Status":
    relatorio = RelatorioStatus(maquina_id, data_inicio, data_fim)
elif tipo_relatorio == "Eficiência":
    relatorio = RelatorioEficiencia(maquina_id, data_inicio, data_fim)
else:
    raise ValueError("Tipo de relatório inválido.")

# Gerar os relatórios nos formatos especificados
relatorio.gerar_relatorios(formatos)
print("Relatórios gerados com sucesso.")
